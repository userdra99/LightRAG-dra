import os
import asyncio
import logging
import numpy as np
from contextlib import asynccontextmanager
import shutil

# --- 1. IMPORTS ---
from fastapi import FastAPI, HTTPException, UploadFile, File
from pydantic import BaseModel

from lightrag.lightrag import LightRAG, QueryParam
from lightrag.utils import EmbeddingFunc
from lightrag.llm.openai import openai_complete_if_cache, openai_embed
from lightrag.kg.shared_storage import initialize_pipeline_status
import pypdf
import docx

# --- 2. SETUP ---
logging.basicConfig(level=logging.INFO)
WORKING_DIR = "./rag_storage"
DATA_DIR = "/mnt/c/Users/User/Downloads/STORIES"
BASE_URL = "http://100.74.68.121:1234/v1"

# --- 3. WRAPPER FUNCTIONS FOR LM STUDIO ---
async def local_llm_func(prompt: str, **kwargs) -> str:
    print(f"...Calling local LLM at {BASE_URL}...")
    return await openai_complete_if_cache(prompt=prompt, model="mistralai/mistral-nemo-instruct-2407", base_url=BASE_URL, api_key="not-needed", **kwargs)

async def local_embedding_func(texts: list[str], **kwargs) -> np.ndarray:
    print(f"...Generating embeddings for {len(texts)} text(s)...")
    return await openai_embed(texts=texts, model="nomic-ai/nomic-embed-text-v1.5", base_url=BASE_URL, api_key="not-needed", **kwargs)

app_state = {}

# --- 4. FASTAPI LIFESPAN MANAGER ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    print("Server starting up...")
    print("Initializing LightRAG engine...")
    rag_instance = LightRAG(
        working_dir=WORKING_DIR,
        llm_model_func=local_llm_func,
        embedding_func=EmbeddingFunc(func=local_embedding_func, embedding_dim=768, max_token_size=8192)
    )
    await rag_instance.initialize_storages()
    await initialize_pipeline_status()
    print("LightRAG engine initialized.")
    
    print(f"Scanning for files in {DATA_DIR} to insert on startup...")
    if os.path.isdir(DATA_DIR):
        for filename in os.listdir(DATA_DIR):
            file_path = os.path.join(DATA_DIR, filename)
            text = ""
            try:
                if filename.endswith(".pdf"): text = "".join(page.extract_text() for page in pypdf.PdfReader(file_path).pages)
                elif filename.endswith(".docx"): text = "\n".join(para.text for para in docx.Document(file_path).paragraphs)
                elif filename.endswith(".txt"):
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f: text = f.read()
                
                if text.strip():
                    await rag_instance.ainsert(input=text, ids=filename)
                    print(f"Successfully processed and inserted: {filename}")
            except Exception as e:
                print(f"Error processing file {filename}: {e}")
    else:
        print(f"Warning: Data directory not found at {DATA_DIR}")

    app_state["rag_instance"] = rag_instance
    print("Server startup complete. Ready to accept queries.")
    
    yield
    
    print("Server shutting down...")
    if app_state.get("rag_instance"):
        await app_state["rag_instance"].finalize_storages()
    print("Shutdown complete.")

app = FastAPI(lifespan=lifespan)

# --- 5. API ENDPOINTS ---
class QueryRequest(BaseModel):
    query_str: str

@app.post("/query")
async def handle_query(request: QueryRequest):
    rag_instance = app_state.get("rag_instance")
    if not rag_instance:
        raise HTTPException(status_code=503, detail="RAG engine is not available")
    
    print(f"Received query: {request.query_str}")
    try:
        # --- NEW 2-STEP LOGIC FOR ROBUST CITATIONS ---

        # STEP 1: Explicitly retrieve ONLY the context (sources) for the query.
        print("Step 1: Retrieving source context for citation...")
        context_response = await rag_instance.aquery(
            query=request.query_str,
            param=QueryParam(mode="mix", top_k=3, only_need_context=True)
        )
        
        sources = []
        if hasattr(context_response, 'context') and context_response.context and hasattr(context_response.context, 'chunks'):
            print("SUCCESS: Found context with chunks. Extracting sources...")
            for chunk in context_response.context.chunks:
                doc_id = getattr(chunk, 'id', 'Unknown')
                sources.append({"text": chunk.text, "doc_id": doc_id})
        else:
             print("Warning: Could not retrieve context for citation.")

        # STEP 2: Generate the final answer using the same query.
        print("Step 2: Generating final answer...")
        answer_response = await rag_instance.aquery(
            query=request.query_str,
            param=QueryParam(mode="mix", top_k=3, only_need_context=False)
        )
        answer = str(answer_response) # The final response is just a string

        return {"answer": answer, "sources": sources}
    except Exception as e:
        logging.error(f"An error occurred during query processing: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error processing the query: {str(e)}")

@app.post("/upload")
async def handle_file_upload(file: UploadFile = File(...)):
    rag_instance = app_state.get("rag_instance")
    if not rag_instance:
        raise HTTPException(status_code=503, detail="RAG engine is not available")
    
    temp_file_path = f"./{file.filename}"
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    text = ""
    print(f"Processing uploaded file: {file.filename}")
    try:
        if file.filename.endswith(".pdf"): text = "".join(page.extract_text() for page in pypdf.PdfReader(temp_file_path).pages)
        elif file.filename.endswith(".docx"): text = "\n".join(para.text for para in docx.Document(temp_file_path).paragraphs)
        elif file.filename.endswith(".txt"):
            with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f: text = f.read()

        if text.strip():
            await rag_instance.ainsert(input=text, ids=file.filename)
            print(f"Successfully inserted content from uploaded file: {file.filename}")
            os.remove(temp_file_path)
            return {"status": "success", "filename": file.filename, "message": "File processed and inserted successfully."}
        else:
            os.remove(temp_file_path)
            raise HTTPException(status_code=400, detail="No text content found in file.")
    except Exception as e:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        logging.error(f"Failed to process uploaded file {file.filename}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to process file: {e}")

@app.get("/")
def read_root():
    return {"status": "LightRAG server is running"}
